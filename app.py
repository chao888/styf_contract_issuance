from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import os
import tempfile

app = Flask(__name__)

class ContractGenerator:
    @staticmethod
    def generate_word_contract(data):
        """生成Word合同文档"""
        doc = Document()
        
        # 设置文档样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
        
        # 标题
        title = doc.add_heading('服务合同', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 合同双方信息
        doc.add_paragraph(f'甲方（委托方）：{data["partyA"]}')
        doc.add_paragraph(f'乙方（受托方）：{data["partyB"]}')
        doc.add_paragraph('')
        
        # 合同条款
        sections = [
            ('第一条 合同标的', f'甲方委托乙方提供以下服务：{data["description"]}。'),
            ('第二条 合同金额', f'本合同总金额为人民币（大写）{ContractGenerator.number_to_chinese(data["amount"])}元整（¥{data["amount"]}）。'),
            ('第三条 合同期限', f'本合同有效期为{data["duration"]}个月，自{data["signDate"]}起至{(datetime.strptime(data["signDate"], "%Y-%m-%d") + timedelta(days=int(data["duration"])*30)).strftime("%Y年%m月%d日")}止。'),
            ('第四条 付款方式', '甲方应按照以下方式向乙方支付合同款项：合同签订后5个工作日内支付合同总金额的50%，服务完成并验收合格后支付剩余50%。'),
            ('第五条 双方权利义务', '''1. 甲方应按照合同约定及时支付合同款项，并提供必要的协助和支持。
2. 乙方应按照合同约定提供优质服务，确保服务质量符合行业标准。
3. 双方应严格履行保密义务，未经对方同意不得向第三方透露合同内容。'''),
            ('第六条 违约责任', '任何一方违反本合同约定，应承担相应的违约责任，并赔偿对方因此遭受的损失。'),
            ('第七条 争议解决', '本合同履行过程中发生的争议，双方应友好协商解决；协商不成的，可向合同签订地人民法院提起诉讼。'),
            ('第八条 其他约定', '''1. 本合同一式两份，甲乙双方各执一份，具有同等法律效力。
2. 本合同自双方签字盖章之日起生效。'''),
            ('第九条 合同签订', f'本合同于{data["signDate"]}在{data["partyA"]}所在地签订。')
        ]
        
        for title, content in sections:
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)
        
        # 签字区域
        doc.add_paragraph('\n\n')
        doc.add_paragraph('甲方（盖章）：_______________')
        doc.add_paragraph('法定代表人/授权代表：_______________')
        doc.add_paragraph('日期：_______________')
        doc.add_paragraph('')
        doc.add_paragraph('乙方（盖章）：_______________')
        doc.add_paragraph('法定代表人/授权代表：_______________')
        doc.add_paragraph('日期：_______________')
        
        return doc
    
    @staticmethod
    def number_to_chinese(num):
        """数字转中文大写"""
        chinese_num = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
        chinese_unit = ['', '拾', '佰', '仟', '万', '拾', '佰', '仟', '亿']
        
        num_str = str(num)
        result = ''
        
        for i in range(len(num_str)):
            n = int(num_str[i])
            unit = chinese_unit[len(num_str) - 1 - i]
            result += chinese_num[n] + unit
        
        return result.replace('零[拾佰仟]', '零').replace('零+万', '万').replace('零+亿', '亿').replace('亿万', '亿').replace('零+', '零').replace('零$', '')

@app.route('/generate_contract', methods=['POST'])
def generate_contract():
    """生成合同API"""
    try:
        data = request.json
        
        # 生成Word文档
        doc = ContractGenerator.generate_word_contract(data)
        
        # 创建临时文件
        temp_dir = tempfile.mkdtemp()
        word_path = os.path.join(temp_dir, '合同.docx')
        doc.save(word_path)
        
        return send_file(word_path, as_attachment=True, download_name='合同.docx')
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/')
def index():
    """主页"""
    return app.send_static_file('index.html')

if __name__ == '__main__':
    app.run(debug=True, port=5000)